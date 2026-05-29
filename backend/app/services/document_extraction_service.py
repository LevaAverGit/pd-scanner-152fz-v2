"""
document_extraction_service.py — policy document detection and text extraction.

Supports text-based PDF (via PyMuPDF) and DOCX (via python-docx).
Image-only / encrypted / unextractable documents are recorded as "unreadable".
No OCR is performed.

Safe boundaries:
- Documents are fetched only from URLs already validated by the URL validator.
- Downloads are capped at MAX_DOCUMENT_BYTES (10 MB).
- No macro execution.  No eval of document content.
- Text is returned as plain strings for keyword analysis only.
"""
from __future__ import annotations

import io
import logging
from urllib.parse import urlparse

import httpx

logger = logging.getLogger(__name__)

MAX_DOCUMENT_BYTES: int = 10 * 1024 * 1024  # 10 MB
DOWNLOAD_TIMEOUT: float = 30.0


# ---------------------------------------------------------------------------
# Type detection
# ---------------------------------------------------------------------------

def detect_document_type(url: str, content_type: str | None = None) -> str:
    """
    Determine document type from URL path extension and optional Content-Type header.

    Returns: "pdf" | "docx" | "doc" | "html"
    - "doc"  → legacy Word binary format — treated as unsupported by the parser.
    - "html" → default when the type cannot be determined.

    URL extension is checked first (most reliable); Content-Type is consulted
    only when the extension is absent or ambiguous.
    """
    path = urlparse(url).path.lower()
    if path.endswith(".pdf"):
        return "pdf"
    if path.endswith(".docx"):
        return "docx"
    if path.endswith(".doc"):
        return "doc"

    if content_type:
        ct = content_type.lower()
        if "application/pdf" in ct:
            return "pdf"
        if "openxmlformats-officedocument.wordprocessingml" in ct:
            return "docx"
        if "application/msword" in ct:
            return "doc"

    return "html"


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_pdf_text(content: bytes) -> tuple[str, str]:
    """
    Extract plain text from PDF bytes using PyMuPDF (fitz).

    Returns (text, status):
      status "parsed"     — text extracted successfully.
      status "unreadable" — PDF opened but yielded no extractable text
                            (likely image-only or protected).
      status "failed"     — could not open or process the PDF.

    No OCR is attempted.  fitz is imported lazily so the module loads even
    if pymupdf is somehow missing from the environment (though it must be
    present at runtime for PDF support to work).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF (pymupdf) not installed — PDF extraction unavailable")
        return "", "failed"

    try:
        doc = fitz.open(stream=content, filetype="pdf")
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text("text"))
        doc.close()
        text = "\n".join(parts).strip()
        if not text:
            return "", "unreadable"
        return text, "parsed"
    except Exception as exc:
        logger.warning("PDF text extraction failed: %s", exc)
        return "", "failed"


def extract_docx_text(content: bytes) -> tuple[str, str]:
    """
    Extract plain text from DOCX bytes using python-docx.

    Returns (text, status):
      status "parsed"     — text extracted successfully.
      status "unreadable" — DOCX opened but contained no text content.
      status "failed"     — could not open or process the file.

    Extracts text from paragraphs and table cells.  No macro execution.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed — DOCX extraction unavailable")
        return "", "failed"

    try:
        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        text = "\n".join(parts).strip()
        if not text:
            return "", "unreadable"
        return text, "parsed"
    except Exception as exc:
        logger.warning("DOCX text extraction failed: %s", exc)
        return "", "failed"


# ---------------------------------------------------------------------------
# HTTP download
# ---------------------------------------------------------------------------

async def download_document(url: str) -> tuple[bytes | None, str | None, str]:
    """
    Download a document from *url* with a hard size cap.

    Returns (content_bytes, content_type_header, status):
      status "ok"        — download succeeded, content_bytes is populated.
      status "too_large" — download aborted after MAX_DOCUMENT_BYTES.
      status "failed"    — HTTP error or network exception.

    Only follows redirects within the httpx default policy.
    The caller is responsible for ensuring the URL has already been through
    the standard URL validator before calling this function.
    """
    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=DOWNLOAD_TIMEOUT,
        ) as client:
            async with client.stream(
                "GET",
                url,
                headers={
                    "User-Agent": (
                        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) "
                        "Chrome/120.0.0.0 Safari/537.36"
                    )
                },
            ) as resp:
                if resp.status_code != 200:
                    logger.warning(
                        "document_extraction: HTTP %d for %s", resp.status_code, url
                    )
                    return None, None, "failed"
                content_type = resp.headers.get("content-type", "")
                chunks: list[bytes] = []
                total = 0
                async for chunk in resp.aiter_bytes(chunk_size=65536):
                    total += len(chunk)
                    if total > MAX_DOCUMENT_BYTES:
                        logger.warning(
                            "document_extraction: size limit exceeded for %s", url
                        )
                        return None, content_type, "too_large"
                    chunks.append(chunk)
                return b"".join(chunks), content_type, "ok"
    except Exception as exc:
        logger.warning("document_extraction: download failed for %s: %s", url, exc)
        return None, None, "failed"
